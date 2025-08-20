import io
import os
import base64
from typing import Optional, Dict, Any
import PyPDF2
from docx import Document
import pandas as pd
from PIL import Image
import openai

class FileProcessor:
    """다양한 파일 형식에서 텍스트를 추출하는 클래스"""
    
    def __init__(self):
        # OpenAI 클라이언트 초기화
        self.client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
    
    async def process_file(self, file_content: bytes, file_extension: str) -> Optional[str]:
        """
        파일 내용을 처리하여 텍스트를 추출합니다.
        
        Args:
            file_content: 파일의 바이트 내용
            file_extension: 파일 확장자
            
        Returns:
            추출된 텍스트 또는 None
        """
        try:
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_content)
            elif file_extension in ['.docx']:
                return self._extract_from_docx(file_content)
            elif file_extension in ['.xlsx', '.xls']:
                return self._extract_from_excel(file_content)
            elif file_extension in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.webp', '.gif']:
                return await self._extract_from_image_with_vision(file_content, file_extension)
            else:
                raise ValueError(f"지원되지 않는 파일 형식: {file_extension}")
        except Exception as e:
            print(f"파일 처리 중 오류 발생: {str(e)}")
            return None
    
    async def analyze_image_with_vision(self, image_content: bytes, image_extension: str, prompt: str = "이 이미지를 분석하고 주요 내용을 설명해주세요.", detail: str = "auto") -> Dict[str, Any]:
        """
        OpenAI Vision API를 사용하여 이미지를 분석합니다.
        
        Args:
            image_content: 이미지 파일의 바이트 내용
            image_extension: 이미지 파일 확장자
            prompt: 분석 요청 프롬프트
            detail: 이미지 분석 상세도 (low, high, auto)
            
        Returns:
            OpenAI API 분석 결과
        """
        try:
            # 이미지를 Base64로 인코딩
            base64_image = base64.b64encode(image_content).decode('utf-8')
            
            # OpenAI Vision API 호출 (chat.completions 사용)
            response = self.client.chat.completions.create(
                model="gpt-4o",  # Vision API 지원 모델
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": prompt
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/{image_extension[1:]};base64,{base64_image}",
                                    "detail": detail
                                }
                            }
                        ]
                    }
                ],
                max_tokens=4000
            )
            
            return {
                "success": True,
                "output_text": response.choices[0].message.content,
                "model": response.model,
                "usage": response.usage.dict() if response.usage else None
            }
            
        except Exception as e:
            print(f"OpenAI Vision API 이미지 분석 오류: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def generate_image_with_gpt(self, prompt: str, size: str = "1024x1024", quality: str = "standard") -> Dict[str, Any]:
        """
        GPT Image 1을 사용하여 이미지를 생성합니다.
        
        Args:
            prompt: 이미지 생성 프롬프트
            size: 이미지 크기 (1024x1024, 1792x1024, 1024x1792)
            quality: 이미지 품질 (standard, hd)
            
        Returns:
            OpenAI API 이미지 생성 결과
        """
        try:
            # OpenAI Image API 호출
            response = self.client.images.generate(
                model="gpt-image-1",
                prompt=prompt,
                size=size,
                quality=quality,
                n=1
            )
            
            return {
                "success": True,
                "image_url": response.data[0].url,
                "revised_prompt": response.data[0].revised_prompt,
                "model": "gpt-image-1"
            }
            
        except Exception as e:
            print(f"OpenAI Image API 이미지 생성 오류: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def analyze_image_with_file_id(self, file_id: str, prompt: str = "이 이미지를 분석해주세요.", detail: str = "auto") -> Dict[str, Any]:
        """
        OpenAI Files API에 업로드된 이미지 파일 ID를 사용하여 분석합니다.
        
        Args:
            file_id: OpenAI Files API의 파일 ID
            prompt: 분석 요청 프롬프트
            detail: 이미지 분석 상세도 (low, high, auto)
            
        Returns:
            OpenAI API 분석 결과
        """
        try:
            # OpenAI Vision API 호출 (파일 ID 사용)
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": prompt
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"https://api.openai.com/v1/files/{file_id}/content",
                                    "detail": detail
                                }
                            }
                        ]
                    }
                ],
                max_tokens=4000
            )
            
            return {
                "success": True,
                "output_text": response.choices[0].message.content,
                "model": response.model,
                "usage": response.usage.dict() if response.usage else None
            }
            
        except Exception as e:
            print(f"OpenAI Vision API 이미지 분석 오류: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def process_pdf_with_openai(self, file_content: bytes, filename: str, prompt: str = "이 PDF를 분석하고 주요 내용을 요약해주세요.") -> Dict[str, Any]:
        """
        OpenAI의 PDF 입력 기능을 사용하여 PDF를 처리합니다.
        
        Args:
            file_content: PDF 파일의 바이트 내용
            filename: 파일명
            prompt: 분석 요청 프롬프트
            
        Returns:
            OpenAI API 분석 결과
        """
        try:
            # PDF를 Base64로 인코딩
            base64_pdf = base64.b64encode(file_content).decode('utf-8')
            
            # OpenAI API 호출 (chat.completions 사용)
            response = self.client.chat.completions.create(
                model="gpt-4o",  # PDF 입력을 지원하는 모델
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": prompt
                            },
                            {
                                "type": "file_url",
                                "file_url": {
                                    "url": f"data:application/pdf;base64,{base64_pdf}"
                                }
                            }
                        ]
                    }
                ],
                max_tokens=4000
            )
            
            return {
                "success": True,
                "output_text": response.choices[0].message.content,
                "model": response.model,
                "usage": response.usage.dict() if response.usage else None
            }
            
        except Exception as e:
            print(f"OpenAI PDF 처리 오류: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def upload_file_to_openai(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        파일을 OpenAI Files API에 업로드합니다.
        
        Args:
            file_content: 파일의 바이트 내용
            filename: 파일명
            
        Returns:
            업로드된 파일 정보
        """
        try:
            # 파일을 임시로 저장하여 업로드
            temp_file_path = f"/tmp/{filename}"
            with open(temp_file_path, "wb") as f:
                f.write(file_content)
            
            # OpenAI Files API에 업로드
            with open(temp_file_path, "rb") as f:
                file_obj = self.client.files.create(
                    file=f,
                    purpose="vision"  # 이미지 분석용
                )
            
            # 임시 파일 삭제
            os.remove(temp_file_path)
            
            return {
                "success": True,
                "file_id": file_obj.id,
                "filename": file_obj.filename,
                "purpose": file_obj.purpose,
                "bytes": file_obj.bytes
            }
            
        except Exception as e:
            print(f"파일 업로드 오류: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def process_pdf_with_file_id(self, file_id: str, prompt: str = "이 파일을 분석하고 주요 내용을 요약해주세요.") -> Dict[str, Any]:
        """
        OpenAI Files API에 업로드된 파일 ID를 사용하여 PDF를 처리합니다.
        
        Args:
            file_id: OpenAI Files API의 파일 ID
            prompt: 분석 요청 프롬프트
            
        Returns:
            OpenAI API 분석 결과
        """
        try:
            # OpenAI API 호출 (파일 ID 사용)
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": prompt
                            },
                            {
                                "type": "file_url",
                                "file_url": {
                                    "url": f"https://api.openai.com/v1/files/{file_id}/content"
                                }
                            }
                        ]
                    }
                ],
                max_tokens=4000
            )
            
            return {
                "success": True,
                "output_text": response.choices[0].message.content,
                "model": response.model,
                "usage": response.usage.dict() if response.usage else None
            }
            
        except Exception as e:
            print(f"OpenAI PDF 처리 오류: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def _extract_from_pdf(self, file_content: bytes) -> str:
        """PDF 파일에서 텍스트 추출"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
        except Exception as e:
            print(f"PDF 텍스트 추출 오류: {str(e)}")
            return ""
    
    def _extract_from_docx(self, file_content: bytes) -> str:
        """DOCX 파일에서 텍스트 추출"""
        try:
            doc = Document(io.BytesIO(file_content))
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # 표에서도 텍스트 추출
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            print(f"DOCX 텍스트 추출 오류: {str(e)}")
            return ""
    
    def _extract_from_excel(self, file_content: bytes) -> str:
        """Excel 파일에서 텍스트 추출"""
        try:
            excel_file = io.BytesIO(file_content)
            
            # 모든 시트 읽기
            excel_data = pd.read_excel(excel_file, sheet_name=None)
            
            text = ""
            for sheet_name, df in excel_data.items():
                text += f"=== {sheet_name} ===\n"
                text += df.to_string(index=False) + "\n\n"
            
            return text.strip()
        except Exception as e:
            print(f"Excel 텍스트 추출 오류: {str(e)}")
            return ""
    
    async def _extract_from_image_with_vision(self, file_content: bytes, file_extension: str) -> str:
        """OpenAI Vision API를 사용하여 이미지에서 텍스트 및 표 추출"""
        try:
            # 이미지를 Base64로 인코딩
            base64_image = base64.b64encode(file_content).decode('utf-8')
            
            # Vision API 호출
            response = self.client.chat.completions.create(
                model="gpt-4o",  # Vision API 지원 모델
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": """이 이미지를 분석하여 다음을 수행해주세요:
1. 이미지에 있는 모든 텍스트를 추출
2. 표가 있다면 표의 구조와 내용을 파악
3. 이미지의 주요 내용을 요약
4. 발견된 정보를 구조화된 형태로 정리

응답은 한국어로 작성해주세요."""
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/{file_extension[1:]};base64,{base64_image}",
                                    "detail": "high"  # 고해상도 분석
                                }
                            }
                        ]
                    }
                ],
                max_tokens=4000,
                temperature=0.1
            )
            
            # 응답에서 텍스트 추출
            extracted_text = response.choices[0].message.content
            
            if not extracted_text:
                return "이미지에서 텍스트를 추출할 수 없습니다."
            
            return extracted_text.strip()
            
        except Exception as e:
            print(f"Vision API 이미지 분석 오류: {str(e)}")
            return f"이미지 분석 중 오류가 발생했습니다: {str(e)}"
    
    def _encode_image_to_base64(self, image_path: str) -> str:
        """이미지 파일을 Base64로 인코딩"""
        try:
            with open(image_path, "rb") as image_file:
                return base64.b64encode(image_file.read()).decode('utf-8')
        except Exception as e:
            print(f"이미지 인코딩 오류: {str(e)}")
            return ""
