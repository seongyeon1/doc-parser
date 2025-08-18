import io
import os
import base64
from typing import Optional
import PyPDF2
from docx import Document
import pandas as pd
from PIL import Image
import openai

class FileProcessor:
    """다양한 파일 형식에서 텍스트를 추출하는 클래스"""
    
    def __init__(self):
        # OpenAI 클라이언트 초기화
        self.client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
    
    async def process_file(self, file_content: bytes, file_extension: str) -> Optional[str]:
        """
        파일 내용을 처리하여 텍스트를 추출합니다.
        
        Args:
            file_content: 파일의 바이트 내용
            file_extension: 파일 확장자
            
        Returns:
            추출된 텍스트 또는 None
        """
        try:
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_content)
            elif file_extension in ['.docx']:
                return self._extract_from_docx(file_content)
            elif file_extension in ['.xlsx', '.xls']:
                return self._extract_from_excel(file_content)
            elif file_extension in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.webp', '.gif']:
                return await self._extract_from_image_with_vision(file_content, file_extension)
            else:
                raise ValueError(f"지원되지 않는 파일 형식: {file_extension}")
        except Exception as e:
            print(f"파일 처리 중 오류 발생: {str(e)}")
            return None
    
    def _extract_from_pdf(self, file_content: bytes) -> str:
        """PDF 파일에서 텍스트 추출"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
        except Exception as e:
            print(f"PDF 텍스트 추출 오류: {str(e)}")
            return ""
    
    def _extract_from_docx(self, file_content: bytes) -> str:
        """DOCX 파일에서 텍스트 추출"""
        try:
            doc = Document(io.BytesIO(file_content))
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # 표에서도 텍스트 추출
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            print(f"DOCX 텍스트 추출 오류: {str(e)}")
            return ""
    
    def _extract_from_excel(self, file_content: bytes) -> str:
        """Excel 파일에서 텍스트 추출"""
        try:
            excel_file = io.BytesIO(file_content)
            
            # 모든 시트 읽기
            excel_data = pd.read_excel(excel_file, sheet_name=None)
            
            text = ""
            for sheet_name, df in excel_data.items():
                text += f"=== {sheet_name} ===\n"
                text += df.to_string(index=False) + "\n\n"
            
            return text.strip()
        except Exception as e:
            print(f"Excel 텍스트 추출 오류: {str(e)}")
            return ""
    
    async def _extract_from_image_with_vision(self, file_content: bytes, file_extension: str) -> str:
        """OpenAI Vision API를 사용하여 이미지에서 텍스트 및 표 추출"""
        try:
            # 이미지를 Base64로 인코딩
            base64_image = base64.b64encode(file_content).decode('utf-8')
            
            # Vision API 호출
            response = self.client.chat.completions.create(
                model="gpt-4o",  # Vision API 지원 모델
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": """이 이미지를 분석하여 다음을 수행해주세요:
1. 이미지에 있는 모든 텍스트를 추출
2. 표가 있다면 표의 구조와 내용을 파악
3. 이미지의 주요 내용을 요약
4. 발견된 정보를 구조화된 형태로 정리

응답은 한국어로 작성해주세요."""
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/{file_extension[1:]};base64,{base64_image}"
                                },
                                "detail": "high"  # 고해상도 분석
                            }
                        ]
                    }
                ],
                max_tokens=4000,
                temperature=0.1
            )
            
            # 응답에서 텍스트 추출
            extracted_text = response.choices[0].message.content
            
            if not extracted_text:
                return "이미지에서 텍스트를 추출할 수 없습니다."
            
            return extracted_text.strip()
            
        except Exception as e:
            print(f"Vision API 이미지 분석 오류: {str(e)}")
            return f"이미지 분석 중 오류가 발생했습니다: {str(e)}"
    
    def _encode_image_to_base64(self, image_path: str) -> str:
        """이미지 파일을 Base64로 인코딩"""
        try:
            with open(image_path, "rb") as image_file:
                return base64.b64encode(image_file.read()).decode('utf-8')
        except Exception as e:
            print(f"이미지 인코딩 오류: {str(e)}")
            return ""
